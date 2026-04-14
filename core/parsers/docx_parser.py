from __future__ import annotations

from io import BytesIO

from docx import Document

from core.language_utils import detect_language
from core.schemas import ParsedDocument


def extract_docx(content: bytes, filename: str) -> ParsedDocument:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    text = "\n".join(paragraphs)
    return ParsedDocument(
        filename=filename,
        file_type="docx",
        text=text,
        detected_language=detect_language(text),
        selectable_text=True,
    )

