from __future__ import annotations

from datetime import datetime
from io import BytesIO
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_premium_markdown(payload: dict[str, Any]) -> str:
    result = payload.get("result") or {}
    customer = payload.get("customer") or {}
    package = payload.get("package") or {}
    blocks = result.get("transformed_blocks") or []
    rewrite = result.get("rewrite_preview") or {}
    plan = result.get("job_search_plan") or {}
    scores = {
        "final_ats_score": result.get("final_ats_score", 0),
        "visibility_score": result.get("visibility_score", 0),
        "ats_parse_score": result.get("ats_parse_score", 0),
    }
    lines = [
        "# Optimized Resume Package",
        "",
        f"- Generated: {datetime.now().isoformat(timespec='seconds')}",
        f"- Package: {package.get('name') or 'Premium simulation'}",
        f"- Customer: {customer.get('name') or 'Local customer'}",
        f"- Target decision: {result.get('decision') or ''}",
        f"- Final ATS score: {scores['final_ats_score']}",
        f"- Visibility score: {scores['visibility_score']}",
        "",
        "## Optimized Resume",
        "",
    ]
    headline = rewrite.get("headline") or _block_after(blocks, "headline")
    summary = rewrite.get("summary") or _block_after(blocks, "summary")
    if headline:
        lines.extend([f"# {headline}", ""])
    if summary:
        lines.extend(["## Summary", "", summary, ""])
    bullets = rewrite.get("bullets") or [_safe(block.get("after")) for block in blocks if str(block.get("id", "")).startswith("bullet_")]
    if bullets:
        lines.extend(["## Experience Highlights", ""])
        lines.extend(f"- {bullet}" for bullet in bullets if bullet)
        lines.append("")
    phrases = rewrite.get("search_phrases") or []
    if phrases:
        lines.extend(["## Skills / Recruiter Search Phrases", ""])
        lines.append(", ".join(str(item) for item in phrases if item))
        lines.append("")
    lines.extend(["## Education", "", "_Keep original education section here if it was not extracted automatically._", ""])
    lines.extend(["## Languages", "", "_Keep original language section here if it was not extracted automatically._", ""])
    lines.extend(["---", "", "## Recommendations Appendix", ""])
    for fix in result.get("top_fixes") or result.get("recommendations") or []:
        lines.append(f"- {fix}")
    lines.extend(["", "## ATS Visibility Notes", ""])
    for note in result.get("why_not_found") or (result.get("missing_signals") or {}).get("why_not_found", []):
        lines.append(f"- {note}")
    lines.extend(["", "## Job Search Plan", ""])
    for role in plan.get("target_roles") or []:
        lines.append(f"- Target role: {role}")
    for query in plan.get("linkedin_queries") or []:
        lines.append(f"- LinkedIn query: `{query}`")
    filters = plan.get("filters") or {}
    if filters:
        lines.append(f"- Filters: {', '.join(f'{key}={value}' for key, value in filters.items())}")
    return "\n".join(lines).strip() + "\n"


def build_premium_docx(payload: dict[str, Any]) -> bytes:
    markdown = build_premium_markdown(payload)
    doc = Document()
    sections = markdown.splitlines()
    for line in sections:
        clean = line.strip()
        if not clean:
            continue
        if clean.startswith("# "):
            paragraph = doc.add_heading(clean[2:].strip(), level=0)
            _set_rtl_if_needed(paragraph, clean)
        elif clean.startswith("## "):
            paragraph = doc.add_heading(clean[3:].strip(), level=1)
            _set_rtl_if_needed(paragraph, clean)
        elif clean.startswith("- "):
            paragraph = doc.add_paragraph(clean[2:].strip(), style="List Bullet")
            _set_rtl_if_needed(paragraph, clean)
        elif clean == "---":
            doc.add_paragraph("")
        else:
            paragraph = doc.add_paragraph(clean)
            _set_rtl_if_needed(paragraph, clean)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def safe_export_filename(payload: dict[str, Any], suffix: str) -> str:
    result = payload.get("result") or {}
    role = result.get("verdict") or result.get("decision") or "ats"
    name = re.sub(r"[^A-Za-z0-9_-]+", "_", role).strip("_").lower()[:38] or "ats_export"
    return f"{name}_premium_resume.{suffix}"


def _block_after(blocks: list[dict[str, Any]], block_id: str) -> str:
    for block in blocks:
        if block.get("id") == block_id:
            return _safe(block.get("after"))
    return ""


def _safe(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _set_rtl_if_needed(paragraph, text: str) -> None:
    if re.search(r"[\u0590-\u05ff]", text):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.right_indent = 0
