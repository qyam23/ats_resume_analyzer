from __future__ import annotations

import argparse
from datetime import datetime, timezone
from io import BytesIO
import json
from pathlib import Path
from typing import Any

from docx import Document

from config.settings import BASE_DIR
from core.product_outputs import build_product_outputs
from core.schemas import JobDescriptionData, ResumeStructuredData
from core.services.analyzer import ATSAnalyzerService
from qa_suite.qa_evaluator import evaluate_case
from qa_suite.qa_reporter import write_reports
from qa_suite.qa_schemas import (
    JobDescriptionCase,
    QAActualResult,
    QAMatrixCase,
    QARunSummary,
    ResumeVariant,
)


FIXTURE_DIR = BASE_DIR / "qa_suite" / "fixtures"
OUTPUT_DIR = BASE_DIR / "qa_suite" / "output"


def main() -> None:
    parser = argparse.ArgumentParser(description="Run internal ATS product QA suite.")
    parser.add_argument("--fixtures", type=Path, default=FIXTURE_DIR, help="Directory containing QA fixture JSON files.")
    parser.add_argument("--output", type=Path, default=OUTPUT_DIR, help="Directory for QA output artifacts.")
    parser.add_argument("--limit", type=int, default=0, help="Optional limit for fast smoke runs.")
    parser.add_argument("--case-id", default="", help="Run a single QA case id.")
    args = parser.parse_args()

    summary = run_qa(fixtures_dir=args.fixtures, output_dir=args.output, limit=args.limit, case_id=args.case_id)
    print(f"QA run {summary.run_id}: {summary.passed} pass, {summary.warnings} warning, {summary.failed} fail / {summary.total_cases} cases")
    print(f"Artifacts: {args.output}")


def run_qa(fixtures_dir: Path = FIXTURE_DIR, output_dir: Path = OUTPUT_DIR, limit: int = 0, case_id: str = "") -> QARunSummary:
    run_id = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H-%M-%SZ")
    fixtures = load_fixtures(fixtures_dir)
    service = ATSAnalyzerService()
    cases = fixtures["matrix"]
    if case_id:
        cases = [case for case in cases if case.case_id == case_id]
    if limit:
        cases = cases[:limit]

    evaluations = []
    for case in cases:
        resume = fixtures["resumes"][case.resume_variant_id]
        jd = fixtures["jobs"][case.jd_id]
        actual = analyze_case(service, resume, jd)
        evaluations.append(evaluate_case(case, actual))

    summary = build_summary(run_id, evaluations)
    write_reports(summary, output_dir)
    return summary


def load_fixtures(fixtures_dir: Path) -> dict[str, Any]:
    resume_payload = _read_json(fixtures_dir / "resume_variants.json")
    job_payload = _read_json(fixtures_dir / "live_job_descriptions.json")
    matrix_payload = _read_json(fixtures_dir / "qa_matrix.json")
    resumes = {item.id: item for item in [ResumeVariant.model_validate(raw) for raw in resume_payload.get("variants", [])]}
    jobs = {item.id: item for item in [JobDescriptionCase.model_validate(raw) for raw in job_payload]}
    matrix = [QAMatrixCase.model_validate(raw) for raw in matrix_payload]
    missing_resume = [case.resume_variant_id for case in matrix if case.resume_variant_id not in resumes]
    missing_jobs = [case.jd_id for case in matrix if case.jd_id not in jobs]
    if missing_resume or missing_jobs:
        raise ValueError(f"Invalid QA matrix references. Missing resumes={missing_resume[:3]}, missing jobs={missing_jobs[:3]}")
    return {"resumes": resumes, "jobs": jobs, "matrix": matrix}


def analyze_case(service: ATSAnalyzerService, resume: ResumeVariant, jd: JobDescriptionCase) -> QAActualResult:
    docx_bytes = _resume_variant_to_docx(resume)
    jd_text = _job_case_to_text(jd)
    result = service.analyze(
        docx_bytes=docx_bytes,
        docx_name=f"{resume.id}.docx",
        jd_text=jd_text,
        public_mode=True,
        use_provider_insights=False,
    )
    data = result.model_dump()
    diagnostics = data.get("detailed_diagnostics", {})
    visibility = diagnostics.get("visibility", {})
    scores = data.get("scores", {})
    resume_data = ResumeStructuredData.model_validate(diagnostics.get("docx_structured") or {})
    jd_data = JobDescriptionData.model_validate(diagnostics.get("job_description") or {})
    final_score = float(scores.get("final_ats_score", 0) or 0)
    product_outputs = build_product_outputs(
        resume=resume_data,
        jd=jd_data,
        missing_keywords=diagnostics.get("missing_keywords", []),
        visibility_profile=visibility,
        final_score=final_score,
        premium_unlocked=False,
    )
    return QAActualResult(
        decision=_decision_for_score(final_score),
        final_ats_score=final_score,
        visibility_score=float(visibility.get("visibility_score", final_score) or 0),
        recruiter_match_score=float(visibility.get("recruiter_match_score", 0) or 0),
        ats_parse_score=float(visibility.get("ats_parse_score", scores.get("parseability_score", 0)) or 0),
        missing_keywords=list(diagnostics.get("missing_keywords", []))[:12],
        why_not_found=list(visibility.get("why_not_found", [])),
        top_fixes=list(visibility.get("top_fixes", [])),
        what_ats_sees=list(visibility.get("what_ats_sees", [])),
        what_recruiter_sees=list(visibility.get("what_recruiter_sees", [])),
        career_suggestions=_career_suggestions_for_qa(resume_data, jd_data),
        transformed_blocks_count=len(product_outputs.get("transformed_blocks", [])),
        rewrite_preview_available=bool(product_outputs.get("rewrite_preview", {}).get("summary")),
        premium_mode=str(product_outputs.get("premium", {}).get("mode", "")),
        job_search_plan_available=bool(product_outputs.get("job_search_plan", {}).get("linkedin_queries")),
        raw={
            "scores": scores,
            "visibility": visibility,
            "transformed_blocks": product_outputs.get("transformed_blocks", []),
            "rewrite_preview": product_outputs.get("rewrite_preview", {}),
            "premium": product_outputs.get("premium", {}),
            "job_search_plan": product_outputs.get("job_search_plan", {}),
        },
    )


def build_summary(run_id: str, evaluations: list) -> QARunSummary:
    created_at = datetime.now(timezone.utc).isoformat()
    counts = {"pass": 0, "warning": 0, "fail": 0}
    by_fit: dict[str, dict[str, int]] = {}
    by_language: dict[str, dict[str, int]] = {}
    for evaluation in evaluations:
        counts[evaluation.status] += 1
        _increment_group(by_fit, evaluation.expected_fit_band, evaluation.status)
        language_key = "+".join(evaluation.source_languages) if evaluation.source_languages else "unknown"
        _increment_group(by_language, language_key, evaluation.status)
    return QARunSummary(
        run_id=run_id,
        created_at=created_at,
        total_cases=len(evaluations),
        passed=counts["pass"],
        warnings=counts["warning"],
        failed=counts["fail"],
        by_fit_band=by_fit,
        by_language=by_language,
        cases=evaluations,
    )


def _resume_variant_to_docx(resume: ResumeVariant) -> bytes:
    doc = Document()
    doc.add_heading(resume.target_role_en or resume.id, level=1)
    doc.add_paragraph(resume.core_positioning_en or resume.core_positioning_he)
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(
        "Senior engineering and operations profile focused on leadership, manufacturing systems, process improvement, supplier coordination, "
        "technical execution, automation, quality, throughput, commissioning and measurable operational results."
    )
    doc.add_heading("Core Skills", level=2)
    for keyword in resume.priority_keywords:
        doc.add_paragraph(keyword, style="List Bullet")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph(f"{resume.target_role_en or 'Engineering Leader'} | Industrial and Manufacturing Projects | 2015-2026")
    for bullet in _variant_bullets(resume):
        doc.add_paragraph(bullet, style="List Bullet")
    doc.add_heading("Languages", level=2)
    doc.add_paragraph("Hebrew, English, Russian")
    handle = BytesIO()
    doc.save(handle)
    return handle.getvalue()


def _job_case_to_text(jd: JobDescriptionCase) -> str:
    return "\n".join(
        [
            f"Company: {jd.company}",
            f"Role title: {jd.role}",
            f"Location: {jd.location}",
            f"Seniority: {jd.fit_strength}",
            "Responsibilities:",
            jd.why,
            "Lead cross-functional engineering execution, manage technical priorities, improve operations, coordinate suppliers, own delivery, quality, cost, schedules and stakeholder communication.",
            "Requirements:",
            f"Experience relevant to {jd.role}, engineering leadership, manufacturing or technical operations, process improvement, documentation, implementation, data-driven decisions and measurable outcomes.",
            f"Language profile: {jd.language_profile}",
            f"Source link: {jd.link}",
        ]
    )


def _variant_bullets(resume: ResumeVariant) -> list[str]:
    keywords = ", ".join(resume.priority_keywords[:5])
    return [
        f"Led engineering execution aligned to {resume.target_role_en}, including {keywords}.",
        "Improved production stability, quality, throughput and operational reliability through structured engineering actions.",
        "Coordinated suppliers, stakeholders, operations and technical teams from requirements through implementation.",
        "Translated technical problems into measurable improvements, documentation, process controls and delivery plans.",
    ]


def _career_suggestions_for_qa(resume: ResumeStructuredData, jd: JobDescriptionData) -> list[str]:
    from core.product_outputs import build_job_search_plan

    plan = build_job_search_plan(resume=resume, jd=jd, visibility_profile={}, missing_keywords=[])
    return list(plan.get("target_roles", []))


def _decision_for_score(score: float) -> str:
    if score >= 78:
        return "APPLY"
    if score >= 45:
        return "APPLY WITH FIXES"
    return "LOW VISIBILITY / LOW PRIORITY"


def _increment_group(groups: dict[str, dict[str, int]], key: str, status: str) -> None:
    groups.setdefault(key or "unknown", {"pass": 0, "warning": 0, "fail": 0})
    groups[key or "unknown"][status] += 1


def _read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8-sig"))


if __name__ == "__main__":
    main()
